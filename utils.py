from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
import os

def load_prompt(filepath):
    with open(filepath, "r", encoding="utf-8") as file:
        return file.read()

def fill_prompt(template, data):
    for key, value in data.items():
        placeholder = "{" + key + "}"

        if not value:
            # Hem başlık hem placeholder varsa komple sil
            # Örnek eşleşme: 📜 CERTIFICATIONS\n{certifications}
            template = re.sub(
                rf"(.*?{key.replace('_', ' ')}.*?\n)?\s*{re.escape(placeholder)}\s*\n?",
                "",
                template,
                flags=re.IGNORECASE
            )
        else:
            template = template.replace(placeholder, value)
    return template

def save_docx_with_formatting(content, filename, user_data):
    doc = Document()

    # === Başlık (isim, ortalanmış) ===
    heading = doc.add_heading(user_data["name"], level=0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # === Kişisel Bilgiler Satırı (emojili, ortalanmış) ===
    info_parts = []
    if user_data.get("email"):
        info_parts.append(f"✉️ {user_data['email']}")
    if user_data.get("portfolio"):
        info_parts.append(f"🔗 {user_data['portfolio']}")
    if user_data.get("github"):
        info_parts.append(f"💻 {user_data['github']}")
    if user_data.get("location"):
        info_parts.append(f"📍 {user_data['location']}")

    contact = doc.add_paragraph(" | ".join(info_parts))
    contact.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    contact.style.font.size = Pt(10)

    doc.add_paragraph("")  # boşluk bırak

    # === İçerik (OpenAI'den gelen metin) ===
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    paragraphs = content.strip().split("\n\n")

    for para in paragraphs:
        para = para.strip()
        if not para:
            continue

        p = doc.add_paragraph()
        parts = re.split(r'(\*\*.*?\*\*)', para)

        for part in parts:
            run = p.add_run()
            if part.startswith("**") and part.endswith("**"):
                run.bold = True
                run.text = part[2:-2]
            else:
                run.text = part

    doc.save(filename)

def save_cv_to_docx(content, filename, user_data):
    save_docx_with_formatting(content, filename, user_data)

def save_cover_letter_to_docx(content, filename, user_data):
    doc = Document()
    doc.add_heading(f"Cover Letter - {user_data['name']}", level=0)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    for line in content.strip().split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    doc.save(filename)

def collect_user_data_interactive():
    print("🔍 Please enter the following details for CV & Cover Letter generation.\n")
    return {
        "name": input("Full Name: "),
        "email": input("Email: "),
        "portfolio": input("Portfolio URL: "),
        "github": input("GitHub URL: "),
        "current_title": input("Current Job Title: "),
        "target_role": input("Target Role: "),
        "industry": input("Industry: "),
        "location": input("Location: "),
        "education_details": input("Education Details: "),
        "experience_details": input("Experience Details: "),
        "technical_skills": input("Technical Skills: "),
        "soft_skills": input("Soft Skills: "),
        "tools": input("Tools / Technologies: "),
        "languages": input("Languages: "),
        "certifications": input("Certifications (if any): "),
        "achievements": input("Achievements (if any): "),
        "projects": input("Projects / Portfolio (if any): "),
        "career_objective": input("Career Objective: "),
        "hobbies": input("Hobbies / Interests: "),
        "company": input("Company you're applying to: "),
        "job_title": input("Job Title you're applying for: "),
        "company_location": input("Company Location: "),
        "job_source": input("Where did you find the job?: "),
        "experience_summary": input("Short Experience Summary: "),
        "key_skills": input("Key Skills: "),
        "motivation": input("What motivates you to apply?: "),
        "mission_alignment": input("How does the company's mission align with yours?: "),
        "personal_connection": input("Any personal connection to the company?: "),
        "tone_preference": input("Preferred tone (e.g., professional, friendly): "),
        "about_keywords": input("Keywords for 'About Me' section (comma-separated): ")
    }
